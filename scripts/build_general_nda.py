from pathlib import Path
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/legal/OSai_General_Mutual_NDA_Draft.md"
OUTPUT = ROOT / "docs/legal/OSai_General_Mutual_NDA_Draft.docx"

NAVY = RGBColor(0x0C, 0x1B, 0x2A)
TEAL = RGBColor(0x00, 0x8F, 0x97)
SLATE = RGBColor(0x55, 0x67, 0x7B)
INK = RGBColor(0x1A, 0x1A, 0x1A)
RED = RGBColor(0x9B, 0x1C, 0x1C)


def font(run, size=10.5, bold=False, italic=False, color=INK, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def field(paragraph, code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = code
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t"); value.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, value, end])
    font(run, 8.5, color=SLATE)


def shade_paragraph(p, fill):
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def add_inline(p, text, size=10.5, color=INK):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        value = part[2:-2] if bold else part
        r = p.add_run(value)
        font(r, size=size, bold=bold, color=color)


def set_repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trpr.append(node)


def new_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum"); abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType"); multi.set(qn("w:val"), "singleLevel"); abstract.append(multi)
    lvl = OxmlElement("w:lvl"); lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
    numfmt = OxmlElement("w:numFmt"); numfmt.set(qn("w:val"), "decimal"); lvl.append(numfmt)
    lvltext = OxmlElement("w:lvlText"); lvltext.set(qn("w:val"), "%1."); lvl.append(lvltext)
    jc = OxmlElement("w:lvlJc"); jc.set(qn("w:val"), "left"); lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "720"); tabs.append(tab); ppr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "720"); ind.set(qn("w:hanging"), "360"); ppr.append(ind)
    spacing = OxmlElement("w:spacing"); spacing.set(qn("w:after"), "100"); spacing.set(qn("w:line"), "264"); spacing.set(qn("w:lineRule"), "auto"); ppr.append(spacing)
    lvl.append(ppr); abstract.append(lvl); numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    aid = OxmlElement("w:abstractNumId"); aid.set(qn("w:val"), str(abstract_id)); num.append(aid); numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId"); numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid]); ppr.append(numpr)


def set_cell(cell, width):
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW"); tcpr.append(tcw)
    tcw.set(qn("w:w"), str(width)); tcw.set(qn("w:type"), "dxa")
    mar = OxmlElement("w:tcMar")
    for name, value in (("top", 100), ("start", 140), ("bottom", 100), ("end", 140)):
        n = OxmlElement(f"w:{name}"); n.set(qn("w:w"), str(value)); n.set(qn("w:type"), "dxa"); mar.append(n)
    tcpr.append(mar)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def signature_table(doc, blocks):
    for idx, lines in enumerate(blocks):
        if idx:
            spacer = doc.add_paragraph(); spacer.paragraph_format.space_after = Pt(2)
        for j, line in enumerate(lines):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            if j == 0:
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.keep_with_next = True
            add_inline(p, line, size=9.5)


def configure(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(1); sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.42); sec.footer_distance = Inches(0.42)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"; normal.font.size = Pt(10.5); normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after, color in (("Heading 1", 15.5, 14, 7, NAVY), ("Heading 2", 12.5, 10, 5, TEAL), ("Heading 3", 11.5, 8, 4, SLATE)):
        st = doc.styles[name]; st.font.name = "Aptos Display"; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after); st.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        st = doc.styles[name]; st.font.name = "Aptos"; st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.5); st.paragraph_format.first_line_indent = Inches(-0.25)
        st.paragraph_format.space_after = Pt(5); st.paragraph_format.line_spacing = 1.10
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("OSai General Mutual NDA | Version 1.0 draft | Page "); font(r, 8.5, color=SLATE); field(fp, "PAGE")


def build():
    doc = Document(); configure(doc)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("GENERAL MUTUAL\nNONDISCLOSURE AGREEMENT"); font(r, 23, bold=True, color=NAVY, name="Aptos Display")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(12)
    r = p.add_run("OSai Work Hub | Version 1.0"); font(r, 13, color=TEAL, name="Aptos Display")
    banner = doc.add_paragraph(); banner.alignment = WD_ALIGN_PARAGRAPH.CENTER; banner.paragraph_format.space_before = Pt(4); banner.paragraph_format.space_after = Pt(12)
    shade_paragraph(banner, "FFF4D6")
    r = banner.add_run("DRAFT FOR BUSINESS-OWNER AND QUALIFIED-COUNSEL REVIEW - NOT FOR SIGNATURE"); font(r, 9, bold=True, color=RED)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_signatures = False
    skip_title = True
    signature_blocks = [[], []]
    sig_index = -1
    current_num_id = None
    for raw in lines:
        line = raw.strip()
        if not line: continue
        if skip_title and line.startswith("# "):
            skip_title = False; continue
        if line.startswith("**Version:") or line.startswith("**Draft date:") or line.startswith("**Status:"):
            continue
        if line == "## Signatures":
            in_signatures = True
            doc.add_paragraph("Signatures", style="Heading 1")
            continue
        if in_signatures:
            if line == "## Pre-production decisions":
                signature_table(doc, signature_blocks)
                p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)
                doc.add_paragraph("Pre-production decisions", style="Heading 1")
                in_signatures = False
                continue
            if line.startswith("### "):
                sig_index += 1
                signature_blocks[sig_index].append(f"**{line[4:]}**")
            elif sig_index >= 0:
                signature_blocks[sig_index].append(line.replace("  ", ""))
            else:
                p = doc.add_paragraph(); add_inline(p, line)
            continue
        if line.startswith("## "):
            current_num_id = None
            doc.add_paragraph(line[3:], style="Heading 1")
        elif re.match(r"^\d+\. ", line):
            if current_num_id is None:
                current_num_id = new_decimal_numbering(doc)
            p = doc.add_paragraph(); apply_numbering(p, current_num_id); add_inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            current_num_id = None
            p = doc.add_paragraph(style="List Bullet"); add_inline(p, line[2:])
        else:
            current_num_id = None
            p = doc.add_paragraph(); add_inline(p, line)

    props = doc.core_properties
    props.title = "OSai General Mutual Nondisclosure Agreement - Draft"
    props.subject = "General NDA for the OSai Work Hub"
    props.author = "Orbit Systems / OSai"
    props.keywords = "OSai, NDA, confidentiality, DocuSign, work hub"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
