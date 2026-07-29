from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from pathlib import Path


ROOT = Path('/Users/earlpowery/Development/osai')
OUT = ROOT / 'OSai_Work_Hub_Terms_of_Use_Draft.docx'

NAVY = RGBColor(0x0C, 0x1B, 0x2A)
TEAL = RGBColor(0x00, 0x9E, 0x9D)
SLATE = RGBColor(0x55, 0x67, 0x7B)
PALE = 'F8F9FA'
TEAL_PALE = 'E7F6F6'
AMBER_PALE = 'FFF4D6'
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)


def set_font(run, size=11, bold=False, italic=False, color=BLACK, name='Aptos'):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def cell_margins(cell, top=110, start=150, bottom=110, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn('w:w'), str(widths[idx]))
            tc_w.set(qn('w:type'), 'dxa')
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = field_code
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t'); text.text = '1'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, sep, text, end])
    set_font(run, size=9, color=SLATE)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style='Normal')
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead); set_font(r, bold=True)
        r = p.add_run(text[len(bold_lead):]); set_font(r)
    else:
        r = p.add_run(text); set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); set_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text); set_font(r)
    return p


def add_callout(doc, label, text, fill=TEAL_PALE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + ' '); set_font(r, bold=True, color=NAVY)
    r = p.add_run(text); set_font(r, color=NAVY)
    return table


doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = Inches(0.82); sec.bottom_margin = Inches(0.78)
sec.left_margin = Inches(1); sec.right_margin = Inches(1)
sec.header_distance = Inches(0.42); sec.footer_distance = Inches(0.42)

# Styles: standard_business_brief, with OSai brand-color override.
normal = doc.styles['Normal']
normal.font.name = 'Aptos'; normal.font.size = Pt(10.5); normal.font.color.rgb = BLACK
normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for style_name, size, before, after, color in [
    ('Heading 1', 16, 16, 8, NAVY), ('Heading 2', 13, 12, 6, TEAL), ('Heading 3', 11.5, 8, 4, SLATE)
]:
    st = doc.styles[style_name]
    st.font.name = 'Aptos Display'; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
for style_name in ['List Bullet', 'List Number']:
    st = doc.styles[style_name]
    st.font.name = 'Aptos'; st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.5)
    st.paragraph_format.first_line_indent = Inches(-0.25)
    st.paragraph_format.space_after = Pt(5)
    st.paragraph_format.line_spacing = 1.10

# Footer.
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = fp.add_run('Draft for legal review  •  July 28, 2026  •  Page '); set_font(r, size=8.5, color=SLATE)
add_field(fp, 'PAGE')

# Masthead.
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
r = p.add_run('TERMS OF USE'); set_font(r, size=25, bold=True, color=NAVY, name='Aptos Display')
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
r = p.add_run('OSai Work Hub'); set_font(r, size=15, color=TEAL, name='Aptos Display')

meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_width(meta, [1900, 7460])
for i, (label, value) in enumerate([
    ('Document status', 'Draft for business-owner and qualified-counsel review'),
    ('Draft date', 'July 28, 2026'),
    ('Effective date', '[INSERT APPROVED EFFECTIVE DATE]'),
    ('Operator', '[INSERT LEGAL ENTITY NAME], operating the OSai Work Hub'),
]):
    shade(meta.cell(i,0), '0C1B2A'); shade(meta.cell(i,1), PALE)
    p = meta.cell(i,0).paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    rr = p.add_run(label.upper()); set_font(rr, size=8.5, bold=True, color=WHITE)
    p = meta.cell(i,1).paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    rr = p.add_run(value); set_font(rr, size=9.5, color=NAVY)

doc.add_paragraph().paragraph_format.space_after = Pt(0)
add_callout(doc, 'Important:', 'These Terms are a tailored working draft, not legal advice. Do not publish them until the bracketed fields and counsel-review items on the final page are resolved.', AMBER_PALE)

doc.add_heading('1. Agreement to these Terms', level=1)
add_body(doc, 'These Terms of Use (the “Terms”) govern your access to and use of the OSai Work Hub, its public pages, authenticated member areas, protected project rooms, beta programs, feedback tools, communications, files, videos, decks, and related services (collectively, the “Hub”). The Hub is operated by [INSERT LEGAL ENTITY NAME] (“OSai,” “we,” “us,” or “our”).')
add_body(doc, 'By accessing or using the Hub, creating an account, accepting an invitation, or selecting an action that indicates acceptance, you agree to these Terms. If you do not agree, do not access or use the Hub. If you use the Hub for an organization, you represent that you have authority to bind that organization, and “you” includes that organization.')
add_body(doc, 'These Terms apply in addition to any nondisclosure agreement, project-specific agreement, beta agreement, statement of work, or other written agreement between you and OSai. If those documents conflict, the more specific agreement controls for its subject matter.')

doc.add_heading('2. What the Hub is', level=1)
add_body(doc, 'OSai is an innovation and venture-development company that creates original ventures, develops products with clients, and provides business and technology consulting. The Hub supports trusted participation in that work by providing identity, access, agreement, project, content, beta, and feedback experiences.')
add_body(doc, 'Public visitors may see general information and project teasers. Additional areas require a verified account, OSai approval, an active access grant, an executed General NDA, a project-specific agreement, a beta assignment, or an internal role. Access is progressive and content remains restricted unless OSai expressly grants the applicable level.')

doc.add_heading('3. Eligibility and authority', level=1)
add_body(doc, 'You may use the Hub only if you can form a binding contract in your jurisdiction, meet the minimum age approved by OSai, and are not barred from using the Hub under applicable law. The initial Hub may be invitation-only or may require approval after an access request. An invitation or verified email does not guarantee membership or project access.')
add_body(doc, 'The approved minimum age and any geographic restrictions must be stated here before publication: [INSERT AGE ELIGIBILITY AND GEOGRAPHIC LIMITS].')

doc.add_heading('4. Accounts and security', level=1)
add_body(doc, 'You must provide accurate, current information; maintain only accounts you are authorized to use; and promptly update material changes. Accounts are personal unless OSai expressly authorizes an organizational or shared account.')
add_bullet(doc, 'Keep credentials, recovery methods, signing links, and access links confidential.')
add_bullet(doc, 'Use reasonable security safeguards and promptly report suspected loss, compromise, or unauthorized use.')
add_bullet(doc, 'Do not allow another person to use your identity, account, invitation, or agreement status.')
add_bullet(doc, 'You are responsible for activity under your account to the extent permitted by law, except activity caused by OSai’s breach of its obligations.')
add_body(doc, 'Authentication confirms identity; it does not by itself grant membership, agreement clearance, project access, beta access, or permission to view protected content.')

doc.add_heading('5. Access decisions and agreements', level=1)
add_body(doc, 'OSai may approve, decline, limit, condition, suspend, expire, or revoke access based on its access policies, project needs, security concerns, agreement status, beta status, or these Terms. Access grants may be project-specific, time-limited, or subject to additional conditions.')
add_body(doc, 'Certain content requires you to review and electronically sign a General NDA or project-specific agreement through DocuSign or another approved provider. You consent to electronic records and signatures for Hub-related agreements and acknowledge that provider status may need to be verified before access is granted. Your separate agreement—not these Terms—defines the scope of confidentiality and other obligations it covers.')

doc.add_heading('6. Confidential and protected materials', level=1)
add_body(doc, 'Protected materials may include concepts, product plans, prototypes, source materials, videos, decks, designs, research, financial or funding information, milestones, beta builds, feedback, and other nonpublic information. You may access and use protected materials only for the purpose and duration authorized by OSai and any applicable agreement.')
add_body(doc, 'Unless OSai expressly authorizes it in writing, you must not:')
add_bullet(doc, 'share, publish, forward, copy, record, download, scrape, or redistribute protected materials;')
add_bullet(doc, 'permit access by an unauthorized person or use another person’s access;')
add_bullet(doc, 'remove confidentiality, copyright, watermark, attribution, or access-control notices; or')
add_bullet(doc, 'use protected materials to compete unfairly, reverse engineer a project, or build or train a product, model, or dataset.')
add_body(doc, 'If a separate confidentiality agreement applies, its terms control. Revocation or expiration of Hub access does not end obligations that survive under that agreement or applicable law.')

doc.add_heading('7. Acceptable use', level=1)
add_body(doc, 'You may use the Hub only for lawful, authorized purposes. You must not:')
for item in [
    'probe, scan, bypass, disable, or interfere with authentication, authorization, rate limits, logging, watermarks, signed links, or other security controls;',
    'access or attempt to access another user’s account, nonpublic area, file, video, project, API, or data without authorization;',
    'introduce malware, destructive code, excessive traffic, automated extraction, or other harmful or disruptive activity;',
    'misrepresent your identity, affiliation, authority, project role, agreement status, beta status, or the source of content;',
    'upload or submit content that is unlawful, infringing, deceptive, abusive, defamatory, harassing, privacy-invasive, or contains sensitive personal data you lack authority to provide;',
    'use Hub content, feedback, or metadata to train artificial-intelligence systems without OSai’s prior written permission; or',
    'use the Hub to market, solicit, offer, sell, or facilitate securities or other regulated products unless OSai has expressly enabled a counsel-approved, compliant process.'
]: add_bullet(doc, item)

doc.add_heading('8. OSai content and intellectual property', level=1)
add_body(doc, 'The Hub and materials supplied through it—including software, text, graphics, brands, designs, audiovisual content, project materials, and compilations—are owned by OSai or its licensors and are protected by intellectual-property and other laws. Except for the limited right to use the Hub as authorized under these Terms, no right, title, or license is granted by implication, estoppel, or otherwise.')
add_body(doc, '“OSai,” “Orbit Systems,” related names, logos, and product names may be trademarks of their respective owners. You may not use them in a way that suggests endorsement, partnership, or affiliation without prior written permission.')

doc.add_heading('9. Your submissions and feedback', level=1)
add_body(doc, 'You retain ownership of content you submit, subject to the rights you grant below and any separate agreement. You represent that you have the rights and permissions needed to submit it and that doing so does not violate law, confidentiality duties, privacy rights, or third-party rights.')
add_body(doc, 'For account, access-request, project, beta, support, and feedback submissions, you grant OSai a nonexclusive, worldwide, royalty-free license to host, store, reproduce, transmit, display, analyze, and otherwise use the submission as reasonably necessary to operate, secure, improve, and administer the Hub and applicable OSai projects. This license ends when the purpose and applicable retention period end, except for lawful backups, audit records, de-identified information, and rights that must survive to honor a separate agreement.')
add_body(doc, 'If you provide ideas, suggestions, bug reports, or other product feedback, OSai may use them without payment or attribution, unless a separate written agreement says otherwise. OSai may use automated or artificial-intelligence tools to help categorize or summarize feedback where disclosed; source submissions should remain traceable and material decisions remain subject to human review.')

doc.add_heading('10. Beta programs and pre-release features', level=1)
add_body(doc, 'Beta builds, prototypes, experiments, and pre-release features may be incomplete, unstable, insecure, changed, suspended, or discontinued without notice. They may contain errors or cause loss or corruption of data. Use them only for authorized evaluation, follow instructions, maintain appropriate backups, and do not rely on them for production, safety-critical, or legally required activity.')
add_body(doc, 'Beta participation may require additional terms, confidentiality obligations, technical-data collection, screenshot consent, release-specific notices, or a project assignment. You may withdraw from a beta as described in the applicable program materials, but prior lawful processing and surviving obligations remain unaffected.')

doc.add_heading('11. Privacy, analytics, and communications', level=1)
add_body(doc, 'Our Privacy Notice explains how we collect, use, retain, disclose, and protect personal information. The Privacy Notice is incorporated by reference to the extent permitted by law: [INSERT PRIVACY NOTICE URL]. If these Terms conflict with the Privacy Notice about personal-information practices, the Privacy Notice controls for that subject.')
add_body(doc, 'The Hub may record security, agreement, access, content, beta, feedback, and audit events. Optional screenshots, device data, session context, or similar information should be collected only with appropriate notice and consent. Legal, security, and access notices are operational communications and may be sent regardless of marketing preferences where permitted by law. You may opt out of nonessential marketing or beta reminders using the available controls.')

doc.add_heading('12. Third-party services and links', level=1)
add_body(doc, 'The Hub may interoperate with services such as authentication, electronic-signature, hosting, storage, video, analytics, email, and project tools. Third-party services may have their own terms and privacy practices. OSai is not responsible for third-party services it does not control, but this sentence does not limit obligations that cannot legally be excluded. Links do not imply endorsement.')

doc.add_heading('13. Changes, availability, and suspension', level=1)
add_body(doc, 'OSai may change, update, limit, suspend, or discontinue all or part of the Hub. We may also update these Terms. If a change is material, we will provide notice appropriate to the circumstances and, where required, obtain renewed acceptance. The “Effective date” will identify the current version. Continued use after an update takes effect constitutes acceptance only to the extent permitted by law.')
add_body(doc, 'We may suspend or terminate access immediately when reasonably necessary to protect the Hub, users, confidential information, intellectual property, legal compliance, or security; to address a suspected violation; or when an invitation, agreement, project membership, or beta assignment expires or is revoked. Where appropriate, we may provide notice and an opportunity to respond.')

doc.add_heading('14. Disclaimers', level=1)
add_body(doc, 'TO THE MAXIMUM EXTENT PERMITTED BY LAW, THE HUB AND ALL BETA, PROJECT, AND OTHER MATERIALS ARE PROVIDED “AS IS” AND “AS AVAILABLE.” OSAI DISCLAIMS ALL EXPRESS, IMPLIED, AND STATUTORY WARRANTIES, INCLUDING MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE, TITLE, NON-INFRINGEMENT, ACCURACY, AVAILABILITY, SECURITY, AND QUIET ENJOYMENT. OSAI DOES NOT WARRANT THAT THE HUB WILL BE UNINTERRUPTED, ERROR-FREE, OR FREE OF HARMFUL COMPONENTS.')
add_body(doc, 'Hub content is for informational and collaborative purposes and is not legal, tax, accounting, investment, medical, or other regulated professional advice. Project information, milestones, projections, funding information, and pre-release materials may change and must not be treated as guaranteed results.')
add_callout(doc, 'No securities offering.', 'Nothing in the Hub is an offer to sell, a solicitation of an offer to buy, or a recommendation of any security or investment. Any expression of interest is non-binding unless and until a separate, counsel-approved process and definitive documents expressly provide otherwise.', TEAL_PALE)

doc.add_heading('15. Limitation of liability', level=1)
add_body(doc, 'TO THE MAXIMUM EXTENT PERMITTED BY LAW, OSAI AND ITS AFFILIATES, LICENSORS, SERVICE PROVIDERS, OFFICERS, DIRECTORS, EMPLOYEES, AND AGENTS WILL NOT BE LIABLE FOR INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, EXEMPLARY, OR PUNITIVE DAMAGES; LOSS OF PROFITS, REVENUE, DATA, GOODWILL, OR BUSINESS OPPORTUNITY; OR UNAUTHORIZED ACCESS TO OR USE OF DATA, ARISING FROM OR RELATED TO THE HUB, EVEN IF ADVISED OF THE POSSIBILITY.')
add_body(doc, 'TO THE MAXIMUM EXTENT PERMITTED BY LAW, THE AGGREGATE LIABILITY OF THOSE PARTIES FOR ALL CLAIMS ARISING FROM OR RELATED TO THE HUB WILL NOT EXCEED THE GREATER OF (A) THE AMOUNT YOU PAID OSAI SPECIFICALLY FOR USE OF THE HUB IN THE 12 MONTHS BEFORE THE EVENT GIVING RISE TO THE CLAIM OR (B) [INSERT COUNSEL-APPROVED FLOOR]. These limits do not apply where prohibited, including to liability that cannot legally be limited.')

doc.add_heading('16. Indemnification', level=1)
add_body(doc, 'To the extent permitted by law, you will defend, indemnify, and hold harmless OSai and its affiliates, licensors, service providers, officers, directors, employees, and agents from third-party claims, damages, judgments, losses, liabilities, costs, and expenses (including reasonable legal fees) arising from your unlawful or unauthorized use of the Hub, your submissions, your breach of these Terms or another applicable agreement, or your infringement or violation of another person’s rights. This obligation does not apply to the extent a claim results from OSai’s own breach, negligence, or willful misconduct.')

doc.add_heading('17. Governing law and disputes', level=1)
add_body(doc, 'These Terms are governed by the laws of [INSERT STATE/COUNTRY], without regard to conflict-of-laws principles. The courts located in [INSERT COUNTY, STATE/COUNTRY] will have exclusive jurisdiction and venue, unless applicable law requires otherwise. [COUNSEL: DECIDE WHETHER TO ADD INFORMAL RESOLUTION, ARBITRATION, CLASS-ACTION WAIVER, SMALL-CLAIMS, AND OPT-OUT TERMS.]')

doc.add_heading('18. General terms', level=1)
add_body(doc, 'You may not assign these Terms or transfer Hub access without OSai’s prior written consent. OSai may assign these Terms in connection with a reorganization, merger, acquisition, sale of assets, or operation of the Hub, subject to applicable law. Failure to enforce a provision is not a waiver. If a provision is unenforceable, it will be limited to the minimum extent necessary and the remainder will remain effective.')
add_body(doc, 'These Terms, the Privacy Notice, and any applicable separate agreements constitute the agreement between you and OSai concerning the Hub. Headings are for convenience only. Provisions that by their nature should survive—including confidentiality, intellectual property, feedback rights, disclaimers, liability limits, indemnification, dispute terms, and audit-related rights—survive termination.')

doc.add_heading('19. Contact and legal notices', level=1)
add_body(doc, 'Questions about these Terms or legal notices should be sent to:')
for line in ['[INSERT LEGAL ENTITY NAME]', '[INSERT POSTAL ADDRESS]', '[INSERT LEGAL/TERMS EMAIL ADDRESS]', '[INSERT WEBSITE URL]']:
    add_bullet(doc, line)

# Counsel checklist starts on a clean page.
p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)
doc.add_heading('Publication readiness checklist', level=1)
add_body(doc, 'This page is an internal drafting aid and should be removed from the public-facing version after the decisions below are completed.')
check_items = [
    ('Legal identity', 'Confirm the legal entity that operates the Hub, owns relevant projects/content, and signs agreements; confirm how “Orbit Systems” and “OSai” relate legally.'),
    ('Effective date/versioning', 'Approve the effective date, notice method for changes, archive/version policy, and renewed-acceptance triggers.'),
    ('Eligibility', 'Set minimum age, invite/access policy, geographic limits, organizational-use rules, and any parental-consent path.'),
    ('Privacy and retention', 'Publish the Privacy Notice; approve retention/deletion rules, cookies/analytics, AI-assisted feedback analysis, screenshots, device/session context, and data-subject request handling.'),
    ('Agreements', 'Confirm precedence among these Terms, the General NDA, project NDAs, beta terms, client agreements, and DocuSign electronic-signature disclosures.'),
    ('IP and feedback', 'Approve the submission license, feedback treatment, AI-training prohibition, confidential-material controls, and client/project-specific ownership rules.'),
    ('Risk allocation', 'Approve warranty disclaimers, liability cap and floor, indemnity, excluded claims, consumer-law carveouts, and any paid-service treatment.'),
    ('Disputes', 'Choose governing law, venue, notice details, and whether arbitration, class waiver, informal resolution, or small-claims language is appropriate.'),
    ('Beta and security', 'Approve beta risk notices, security-reporting channel, vulnerability-testing policy, incident notice approach, and export/sanctions language if needed.'),
    ('Investment boundary', 'Confirm counsel-approved language and keep investment transactions disabled until a compliant provider and offering structure are approved.'),
    ('Accessibility/consent', 'Confirm how acceptance is recorded, ensure the Terms remain accessible before and after account creation, and provide a printable/downloadable version.'),
    ('Contact details', 'Insert legal name, postal address, legal email, privacy contact, security contact, and canonical URLs.'),
]
for idx, (label, detail) in enumerate(check_items, 1):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(label + ': '); set_font(r, bold=True, color=NAVY)
    r = p.add_run(detail); set_font(r)

doc.add_heading('Recommended acceptance record', level=2)
add_body(doc, 'For each acceptance, retain the immutable user ID, Terms version/effective date, acceptance timestamp, method or interface, relevant account/project context, and an auditable reference to the exact text presented. Do not use email as the relational key.')

# Core properties.
doc.core_properties.title = 'OSai Work Hub Terms of Use — Draft'
doc.core_properties.subject = 'Terms governing access to and use of the OSai Work Hub'
doc.core_properties.author = 'OSai'
doc.core_properties.keywords = 'OSai, Work Hub, Terms of Use, member access, projects, beta'
doc.core_properties.comments = 'Draft for business-owner and qualified-counsel review.'

doc.save(OUT)
print(OUT)
