from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/legal/OSai_Work_Hub_Privacy_Policy_Draft.md"
OUTPUT = ROOT / "docs/legal/OSai_Work_Hub_Privacy_Policy_Draft.docx"

NAVY = RGBColor(18, 41, 68)
BLUE = RGBColor(43, 111, 165)
MUTED = RGBColor(92, 101, 112)
RED = RGBColor(160, 35, 42)
LIGHT_RED = "FCE8E6"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in (
        ("Title", 26, 0, 6, NAVY),
        ("Subtitle", 12, 0, 16, MUTED),
        ("Heading 1", 16, 18, 7, NAVY),
        ("Heading 2", 12.5, 12, 5, BLUE),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name.startswith("Heading")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")

    list_style = styles["List Bullet"]
    list_style.font.name = "Arial"
    list_style.font.size = Pt(10.5)
    list_style.paragraph_format.left_indent = Inches(0.45)
    list_style.paragraph_format.first_line_indent = Inches(-0.2)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.15


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("OSai Work Hub Privacy Policy  |  Draft — July 28, 2026")
    r.font.name = "Arial"
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED


def add_draft_banner(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_RED)
    p_pr.append(shd)
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), "160")
    spacing.set(qn("w:after"), "160")
    r = p.add_run("DRAFT FOR LEGAL AND OPERATIONAL REVIEW — DO NOT PUBLISH")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = RED


def add_markdown_inline(paragraph, text):
    # Handles the limited bold markup used by this controlled source.
    parts = text.split("**")
    for index, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = index % 2 == 1


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(doc)
    add_footer(section)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_checklist = False
    first_title = True
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue
        if line == "---":
            continue
        if line == "**Draft for legal and operational review — do not publish yet**":
            continue
        if line.startswith("# "):
            title = line[2:]
            if first_title:
                p = doc.add_paragraph(style="Title")
                p.paragraph_format.space_after = Pt(3)
                add_markdown_inline(p, title)
                sub = doc.add_paragraph(style="Subtitle")
                sub.add_run("Privacy notice for the public website, member hub, protected project rooms, and beta programs")
                add_draft_banner(doc)
                first_title = False
            else:
                doc.add_section(WD_SECTION.NEW_PAGE)
                in_checklist = True
                p = doc.add_paragraph(title, style="Heading 1")
                p.runs[0].font.color.rgb = RED
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_markdown_inline(p, line[2:])
            continue
        if in_checklist and line[:2].rstrip(".").isdigit():
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(5)
            add_markdown_inline(p, line.split(". ", 1)[1])
            continue
        p = doc.add_paragraph()
        add_markdown_inline(p, line)
        if line.startswith("Before publication"):
            for run in p.runs:
                run.font.color.rgb = RED
                run.italic = True

    props = doc.core_properties
    props.title = "OSai Work Hub Privacy Policy"
    props.subject = "Draft privacy policy for the OSai public website and authenticated work hub"
    props.author = "Orbit Systems / OSai"
    props.keywords = "OSai, privacy policy, member hub, beta, DocuSign, Neon Auth"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
